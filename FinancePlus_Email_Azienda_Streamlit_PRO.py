# -*- coding: utf-8 -*-
"""
FinancePlus Email Azienda Streamlit PRO
Versione 1.0 - Cerca azienda, vedi tutto, scarica tutto, archivia mail e allegati per mittente/azienda.

Avvio:
    streamlit run FinancePlus_Email_Azienda_Streamlit_PRO.py

Funzioni principali:
- CERCA AZIENDA: cerca nelle email per oggetto, contenuto, mittente, nomi allegati e testo estratto dagli allegati.
- SCARICA TUTTO: salva tutte le email e gli allegati riferiti all'azienda.
- VEDI TUTTO: mostra anteprima, sintesi intelligente locale/AI opzionale, allegati e selezione puntuale.
- SCARICA SELEZIONATI: salva solo gli elementi scelti.
- Archivio: MITTENTE / AZIENDA / documenti senza cartelle anno-mese.
- Rinomina allegati con data ricezione mail: nomefile_06_MAGGIO_2026.ext
- Stampa PDF della mail: MAIL_06_MAGGIO_2026.pdf
- Database SQLite con storico scarichi e doppioni SHA-256.

Nota sicurezza:
- Per Gmail usare una App Password, non la password principale.
- Le credenziali non vengono salvate dal programma.
"""

from __future__ import annotations

import base64
import dataclasses
import datetime as dt
import email
import hashlib
import html
import imaplib
import io
import json
import mimetypes
import os
import re
import sqlite3
import textwrap
from collections import Counter, defaultdict
from email.header import decode_header
from email.message import Message
from email.utils import parseaddr, parsedate_to_datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import streamlit as st
import streamlit.components.v1 as components

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    import openpyxl
except Exception:  # pragma: no cover
    openpyxl = None

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
except Exception:  # pragma: no cover
    SimpleDocTemplate = None

try:
    from bs4 import BeautifulSoup
except Exception:  # pragma: no cover
    BeautifulSoup = None


APP_TITLE = "FinancePlus - Cerca Azienda Email PRO"
APP_VERSION = "1.0 Streamlit"
DEFAULT_ARCHIVE = str(Path.home() / "FinancePlus_Email_Aziende_Archivio")

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp"}
PREVIEW_TEXT_EXTENSIONS = {".txt", ".csv", ".log", ".xml", ".json", ".html", ".htm"}
OFFICE_EXTENSIONS = {".docx", ".xlsx"}

MESI_IT = {
    1: "GENNAIO",
    2: "FEBBRAIO",
    3: "MARZO",
    4: "APRILE",
    5: "MAGGIO",
    6: "GIUGNO",
    7: "LUGLIO",
    8: "AGOSTO",
    9: "SETTEMBRE",
    10: "OTTOBRE",
    11: "NOVEMBRE",
    12: "DICEMBRE",
}

DOC_TYPE_KEYWORDS = {
    "BILANCIO": ["bilancio", "stato patrimoniale", "conto economico", "nota integrativa", "attivo", "passivo"],
    "CENTRALE_RISCHI": ["centrale rischi", "banca d'italia", "accordato", "utilizzato", "garanzie", "sofferenze"],
    "VISURA": ["visura", "registro imprese", "rea", "ateco", "camera di commercio", "cciaa"],
    "ESTRATTO_CONTO": ["estratto conto", "saldo iniziale", "saldo finale", "iban", "movimenti", "dare", "avere"],
    "DURC": ["durc", "regolarita contributiva", "inps", "inail", "cnce"],
    "FATTURA": ["fattura", "imponibile", "iva", "totale documento", "sdi", "scadenza pagamento"],
    "CONTRATTO": ["contratto", "accordo", "scrittura privata", "clausola", "corrispettivo", "firma"],
    "DOCUMENTO_IDENTITA": ["carta d'identita", "carta di identita", "patente", "passaporto", "documento di identita"],
    "BUSINESS_PLAN": ["business plan", "piano economico", "piano finanziario", "cash flow", "dscr", "previsionale"],
    "PREVENTIVO": ["preventivo", "offerta", "proposta economica", "validita offerta", "quotazione"],
    "MANDATO": ["mandato", "incarico professionale", "consulenza finanziaria", "compenso professionale"],
}

COMPANY_LEGAL_SUFFIX = r"(?:S\.?\s*R\.?\s*L\.?|SRL|S\.?\s*P\.?\s*A\.?|SPA|S\.?\s*A\.?\s*S\.?|SAS|S\.?\s*N\.?\s*C\.?|SNC|SOCIETA'?\s+COOPERATIVA|COOPERATIVA|CONSORZIO|IMPRESA\s+INDIVIDUALE|DITTA\s+INDIVIDUALE)"
COMPANY_REGEX = re.compile(rf"\b([A-Z0-9][A-Z0-9&'\.\- ]{{2,80}}\s+{COMPANY_LEGAL_SUFFIX})\b", re.IGNORECASE)


@dataclasses.dataclass
class AttachmentData:
    index: int
    filename: str
    content_type: str
    size: int
    sha256: str
    data: bytes
    extracted_text: str = ""
    document_type: str = "ALTRO"


@dataclasses.dataclass
class EmailRecord:
    index: int
    uid: str
    message_id: str
    subject: str
    from_header: str
    sender_name: str
    sender_email: str
    date: dt.datetime
    body_text: str
    attachments: List[AttachmentData]
    detected_company: str
    company_confidence: int
    match_score: int
    summary: str
    raw_size: int = 0


# -----------------------------
# Utility testo, date, nomi file
# -----------------------------

def decode_mime_header(value: Optional[str]) -> str:
    if not value:
        return ""
    out = []
    for fragment, encoding in decode_header(value):
        if isinstance(fragment, bytes):
            try:
                out.append(fragment.decode(encoding or "utf-8", errors="replace"))
            except Exception:
                out.append(fragment.decode("latin-1", errors="replace"))
        else:
            out.append(fragment)
    return "".join(out).strip()


def html_to_text(raw_html: str) -> str:
    if not raw_html:
        return ""
    if BeautifulSoup is not None:
        try:
            return BeautifulSoup(raw_html, "html.parser").get_text("\n")
        except Exception:
            pass
    txt = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw_html)
    txt = re.sub(r"(?s)<br\s*/?>", "\n", txt)
    txt = re.sub(r"(?s)</p>", "\n", txt)
    txt = re.sub(r"(?s)<.*?>", " ", txt)
    return html.unescape(txt)


def clean_spaces(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_for_match(text: str) -> str:
    text = text or ""
    replacements = {
        "à": "a", "è": "e", "é": "e", "ì": "i", "ò": "o", "ù": "u",
        "À": "A", "È": "E", "É": "E", "Ì": "I", "Ò": "O", "Ù": "U",
    }
    for a, b in replacements.items():
        text = text.replace(a, b)
    text = text.upper()
    text = re.sub(r"[^A-Z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def sanitize_filename(value: str, default: str = "SENZA_NOME", max_len: int = 120) -> str:
    value = value or default
    value = decode_mime_header(value)
    value = re.sub(r"[<>:\\|?*\x00-\x1F]", "_", value)
    value = re.sub(r"\s+", " ", value).strip(" ._")
    if not value:
        value = default
    return value[:max_len]


def sanitize_folder(value: str, default: str = "SENZA_NOME") -> str:
    value = sanitize_filename(value, default=default, max_len=90)
    value = value.replace("@", "_AT_")
    return value


def date_label_italian(value: dt.datetime) -> str:
    if value.tzinfo:
        value = value.astimezone().replace(tzinfo=None)
    return f"{value.day:02d}_{MESI_IT[value.month]}_{value.year}"


def parse_email_date(value: Optional[str]) -> dt.datetime:
    if not value:
        return dt.datetime.now()
    try:
        parsed = parsedate_to_datetime(value)
        if parsed is None:
            return dt.datetime.now()
        if parsed.tzinfo:
            parsed = parsed.astimezone().replace(tzinfo=None)
        return parsed
    except Exception:
        return dt.datetime.now()


def append_mail_date_to_filename(filename: str, mail_date: dt.datetime) -> str:
    safe = sanitize_filename(filename, default="ALLEGATO")
    path = Path(safe)
    suffix = path.suffix.lower()
    stem = path.stem or "ALLEGATO"
    # Evita suffissi data doppi se il file viene riscaricato.
    stem = re.sub(r"_\d{2}_(GENNAIO|FEBBRAIO|MARZO|APRILE|MAGGIO|GIUGNO|LUGLIO|AGOSTO|SETTEMBRE|OTTOBRE|NOVEMBRE|DICEMBRE)_\d{4}$", "", stem, flags=re.I)
    return f"{stem}_{date_label_italian(mail_date)}{suffix}"


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    parent = path.parent
    i = 2
    while True:
        candidate = parent / f"{stem}_{i}{suffix}"
        if not candidate.exists():
            return candidate
        i += 1


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def is_image_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in IMAGE_EXTENSIONS


# -----------------------------
# Estrazione testo allegati
# -----------------------------

def extract_text_from_pdf(data: bytes, max_pages: int = 5) -> str:
    chunks: List[str] = []
    if fitz is not None:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            for page in doc[:max_pages]:
                chunks.append(page.get_text("text"))
            return clean_spaces("\n".join(chunks))[:30000]
        except Exception:
            pass
    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages[:max_pages]:
                    chunks.append(page.extract_text() or "")
            return clean_spaces("\n".join(chunks))[:30000]
        except Exception:
            pass
    return ""


def extract_text_from_docx(data: bytes, max_chars: int = 30000) -> str:
    if Document is None:
        return ""
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables[:10]:
            for row in table.rows[:50]:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return clean_spaces("\n".join(parts))[:max_chars]
    except Exception:
        return ""


def extract_text_from_xlsx(data: bytes, max_rows: int = 120, max_chars: int = 30000) -> str:
    if openpyxl is None:
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        chunks: List[str] = []
        for ws in wb.worksheets[:5]:
            chunks.append(f"Foglio: {ws.title}")
            for r, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if r > max_rows:
                    break
                vals = [str(v) for v in row if v is not None]
                if vals:
                    chunks.append(" | ".join(vals))
        return clean_spaces("\n".join(chunks))[:max_chars]
    except Exception:
        return ""


def extract_text_from_attachment(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(data)
    if ext == ".docx":
        return extract_text_from_docx(data)
    if ext == ".xlsx":
        return extract_text_from_xlsx(data)
    if ext in PREVIEW_TEXT_EXTENSIONS:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return clean_spaces(data.decode(enc, errors="replace"))[:30000]
            except Exception:
                pass
    return ""


# -----------------------------
# Classificazione e analisi azienda
# -----------------------------

def classify_document_type(filename: str, text: str = "") -> str:
    source = normalize_for_match(f"{filename}\n{text[:10000]}")
    best_type = "ALTRO"
    best_score = 0
    for doc_type, words in DOC_TYPE_KEYWORDS.items():
        score = 0
        for word in words:
            if normalize_for_match(word) in source:
                score += 2 if len(word.split()) > 1 else 1
        if score > best_score:
            best_type = doc_type
            best_score = score
    return best_type if best_score > 0 else "ALTRO"


def find_company_candidates(text: str, weight: int = 10) -> Counter:
    candidates: Counter = Counter()
    for match in COMPANY_REGEX.finditer(text or ""):
        company = re.sub(r"\s+", " ", match.group(1)).strip(" ,.;:-")
        company = normalize_company_display(company)
        if len(company) >= 4:
            candidates[company] += weight
    return candidates


def normalize_company_display(value: str) -> str:
    value = clean_spaces(value or "")
    value = value.strip(" ,.;:-_/\\")
    value = re.sub(r"\bS\s*\.?\s*R\s*\.?\s*L\s*\.?\b", "S.R.L.", value, flags=re.I)
    value = re.sub(r"\bS\s*\.?\s*P\s*\.?\s*A\s*\.?\b", "S.P.A.", value, flags=re.I)
    value = re.sub(r"\bS\s*\.?\s*A\s*\.?\s*S\s*\.?\b", "S.A.S.", value, flags=re.I)
    value = re.sub(r"\bS\s*\.?\s*N\s*\.?\s*C\s*\.?\b", "S.N.C.", value, flags=re.I)
    # Togli prefissi numerici tipici di data/progressivi: 07 SCHIANO SRL -> SCHIANO SRL
    value = re.sub(r"^\d{1,3}\s+", "", value).strip()
    return value.upper()


def infer_company_from_email(subject: str, body: str, attachment_names: Iterable[str], attachment_texts: Iterable[str], query_company: str) -> Tuple[str, int, Dict[str, int]]:
    """Rileva azienda usando oggetto, corpo mail, nomi allegati e testo estratto dagli allegati.

    Regola prioritaria richiesta dall'utente:
    - Se l'oggetto contiene "documentazione BelGarden", la mail va in BelGarden.
    - Quindi oggetto e contenuto mail pesano molto piu' delle controparti casuali.
    """
    candidates: Counter = Counter()
    query_clean = normalize_company_display(query_company) if query_company else ""
    subject_norm = normalize_for_match(subject)
    body_norm = normalize_for_match(body[:15000])
    names_text = "\n".join(attachment_names)
    names_norm = normalize_for_match(names_text)
    attach_text = "\n".join(t[:8000] for t in attachment_texts if t)
    attach_norm = normalize_for_match(attach_text)

    if query_clean:
        qn = normalize_for_match(query_clean)
        if qn and qn in subject_norm:
            candidates[query_clean] += 90
        if qn and qn in body_norm:
            candidates[query_clean] += 70
        if qn and qn in names_norm:
            candidates[query_clean] += 60
        if qn and qn in attach_norm:
            candidates[query_clean] += 50
        # Se l'utente sta cercando quell'azienda, mantienila candidata anche se il testo ha piccole varianti.
        if any(token in subject_norm or token in body_norm or token in names_norm or token in attach_norm for token in qn.split() if len(token) >= 4):
            candidates[query_clean] += 25

    # Pattern con forma giuridica. Oggetto pesa piu' del corpo, corpo piu' degli allegati.
    candidates.update(find_company_candidates(subject, weight=80))
    candidates.update(find_company_candidates(body[:5000], weight=55))
    candidates.update(find_company_candidates(names_text, weight=45))
    candidates.update(find_company_candidates(attach_text[:12000], weight=35))

    # Pattern semantici: documentazione/per/azienda/cliente/societa X.
    semantic_source = f"{subject}\n{body[:3000]}"
    semantic_patterns = [
        r"(?:documentazione|documenti|pratica|cliente|azienda|societa|società|impresa|bilancio|visura|contratto|preventivo)\s+(?:per|di|della|del|relativa a|inerente a)?\s*([A-Z0-9][A-Z0-9&'\.\- ]{3,50})",
        r"(?:riferit[ao] a|per conto di|in merito a)\s+([A-Z0-9][A-Z0-9&'\.\- ]{3,50})",
    ]
    for pat in semantic_patterns:
        for m in re.finditer(pat, semantic_source, flags=re.I):
            raw = m.group(1)
            raw = re.split(r"[\n\r,;:\.]", raw)[0]
            raw = normalize_company_display(raw)
            # Evita parole troppo generiche.
            if 3 < len(raw) <= 70 and raw not in {"DOCUMENTAZIONE", "AZIENDA", "CLIENTE", "SOCIETA"}:
                if query_clean and normalize_for_match(query_clean) in normalize_for_match(raw + " " + query_clean):
                    candidates[query_clean] += 45
                else:
                    candidates[raw] += 30

    if not candidates:
        return (query_clean or "DA_VEDERE", 0, {})

    # Penalizza candidati chiaramente generici o troppo lunghi.
    adjusted: Dict[str, int] = {}
    generic = {"BANCA", "INTESA", "UNICREDIT", "BPER", "BANCO BPM", "DOCUMENTAZIONE", "CLIENTE", "AZIENDA"}
    for c, score in candidates.items():
        cn = normalize_for_match(c)
        if cn in generic or len(cn) < 3:
            continue
        if len(c.split()) > 10:
            score -= 20
        adjusted[c] = max(0, int(score))

    if not adjusted:
        return (query_clean or "DA_VEDERE", 0, {})

    best, best_score = max(adjusted.items(), key=lambda kv: kv[1])
    confidence = min(100, max(0, best_score))
    return best, confidence, adjusted


def email_match_score(subject: str, body: str, from_header: str, attachment_names: Iterable[str], attachment_texts: Iterable[str], query: str) -> int:
    if not query:
        return 0
    q = normalize_for_match(query)
    if not q:
        return 0
    score = 0
    subject_n = normalize_for_match(subject)
    body_n = normalize_for_match(body[:15000])
    from_n = normalize_for_match(from_header)
    att_names_n = normalize_for_match(" ".join(attachment_names))
    att_text_n = normalize_for_match("\n".join(t[:8000] for t in attachment_texts if t))
    if q in subject_n:
        score += 70
    if q in body_n:
        score += 55
    if q in att_names_n:
        score += 40
    if q in att_text_n:
        score += 45
    if q in from_n:
        score += 20
    tokens = [t for t in q.split() if len(t) >= 4]
    for t in tokens:
        if t in subject_n:
            score += 10
        if t in body_n:
            score += 7
        if t in att_names_n:
            score += 7
        if t in att_text_n:
            score += 5
    return min(100, score)


def local_summary(record: EmailRecord) -> str:
    doc_types = Counter([a.document_type for a in record.attachments])
    attachments = ", ".join(a.filename for a in record.attachments[:8]) or "nessun allegato"
    body_excerpt = clean_spaces(record.body_text)[:900]
    if len(record.body_text) > 900:
        body_excerpt += "..."
    doc_type_txt = ", ".join(f"{k}({v})" for k, v in doc_types.items()) or "nessuno"
    return (
        f"Azienda stimata: {record.detected_company} - confidenza {record.company_confidence}%.\n"
        f"Tipologie allegati: {doc_type_txt}.\n"
        f"Allegati: {attachments}.\n"
        f"Sintesi contenuto: {body_excerpt if body_excerpt else 'Corpo email non disponibile o vuoto.'}"
    )


def ai_summary_optional(record: EmailRecord, api_key: str, model: str) -> str:
    """Sintesi AI opzionale. Se non configurata, usa sintesi locale."""
    if not api_key:
        return local_summary(record)
    try:
        from openai import OpenAI  # type: ignore
        client = OpenAI(api_key=api_key)
        prompt = f"""
Sei un assistente documentale per FinancePlus. Sintetizza in italiano questa email in modo operativo.
Indica: azienda, oggetto, motivo della mail, allegati, dati utili per archiviazione, eventuali criticita'.

Mittente: {record.from_header}
Data: {record.date}
Oggetto: {record.subject}
Azienda stimata: {record.detected_company}
Allegati: {', '.join(a.filename for a in record.attachments)}
Corpo email:
{record.body_text[:12000]}
"""
        response = client.chat.completions.create(
            model=model or "gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
        )
        return response.choices[0].message.content or local_summary(record)
    except Exception as exc:
        return local_summary(record) + f"\n\nNota: sintesi AI non disponibile ({exc})."


# -----------------------------
# IMAP e parsing email
# -----------------------------

def connect_imap(host: str, port: int, username: str, password: str, mailbox: str, use_ssl: bool = True) -> imaplib.IMAP4:
    if use_ssl:
        mail = imaplib.IMAP4_SSL(host, int(port))
    else:
        mail = imaplib.IMAP4(host, int(port))
    mail.login(username, password)
    status, _ = mail.select(mailbox or "INBOX")
    if status != "OK":
        raise RuntimeError(f"Impossibile aprire la cartella email: {mailbox}")
    return mail


def make_since_criterion(days_back: int) -> str:
    since = dt.datetime.now() - dt.timedelta(days=max(1, int(days_back)))
    return since.strftime("%d-%b-%Y")


def fetch_email_ids(mail: imaplib.IMAP4, days_back: int, max_emails: int) -> List[bytes]:
    criterion = f'(SINCE "{make_since_criterion(days_back)}")'
    status, data = mail.search(None, criterion)
    if status != "OK" or not data or not data[0]:
        status, data = mail.search(None, "ALL")
    ids = data[0].split() if data and data[0] else []
    ids = list(reversed(ids))
    return ids[:max(1, int(max_emails))]


def parse_email_message(index: int, uid: str, raw: bytes, query_company: str) -> EmailRecord:
    msg = email.message_from_bytes(raw)
    subject = decode_mime_header(msg.get("Subject")) or "Senza oggetto"
    from_header = decode_mime_header(msg.get("From"))
    sender_name, sender_email = parseaddr(from_header)
    sender_name = decode_mime_header(sender_name) or sender_email or "Mittente_sconosciuto"
    date_value = parse_email_date(msg.get("Date"))
    message_id = msg.get("Message-ID", uid)

    body_parts: List[str] = []
    html_parts: List[str] = []
    attachments: List[AttachmentData] = []

    if msg.is_multipart():
        walk_parts = msg.walk()
    else:
        walk_parts = [msg]

    att_index = 0
    for part in walk_parts:
        content_disposition = (part.get("Content-Disposition") or "").lower()
        content_type = part.get_content_type() or "application/octet-stream"
        filename = part.get_filename()
        if filename:
            filename = decode_mime_header(filename)
            data = part.get_payload(decode=True) or b""
            att_text = extract_text_from_attachment(filename, data)
            doc_type = classify_document_type(filename, att_text)
            attachments.append(
                AttachmentData(
                    index=att_index,
                    filename=filename,
                    content_type=content_type,
                    size=len(data),
                    sha256=sha256_bytes(data),
                    data=data,
                    extracted_text=att_text,
                    document_type=doc_type,
                )
            )
            att_index += 1
            continue

        if "attachment" in content_disposition:
            continue

        payload = part.get_payload(decode=True)
        charset = part.get_content_charset() or "utf-8"
        if payload is None:
            continue
        try:
            text = payload.decode(charset, errors="replace")
        except Exception:
            text = payload.decode("latin-1", errors="replace")
        if content_type == "text/plain":
            body_parts.append(text)
        elif content_type == "text/html":
            html_parts.append(html_to_text(text))

    body_text = clean_spaces("\n".join(body_parts) or "\n".join(html_parts))
    attachment_names = [a.filename for a in attachments]
    attachment_texts = [a.extracted_text for a in attachments]
    company, confidence, _scores = infer_company_from_email(subject, body_text, attachment_names, attachment_texts, query_company)
    match = email_match_score(subject, body_text, from_header, attachment_names, attachment_texts, query_company)

    temp_record = EmailRecord(
        index=index,
        uid=uid,
        message_id=message_id,
        subject=subject,
        from_header=from_header,
        sender_name=sender_name,
        sender_email=sender_email,
        date=date_value,
        body_text=body_text,
        attachments=attachments,
        detected_company=company,
        company_confidence=confidence,
        match_score=match,
        summary="",
        raw_size=len(raw),
    )
    temp_record.summary = local_summary(temp_record)
    return temp_record


def search_company_emails(
    host: str,
    port: int,
    username: str,
    password: str,
    mailbox: str,
    query_company: str,
    days_back: int,
    max_emails: int,
    min_score: int,
    use_ssl: bool = True,
) -> List[EmailRecord]:
    mail = connect_imap(host, port, username, password, mailbox, use_ssl=use_ssl)
    try:
        ids = fetch_email_ids(mail, days_back=days_back, max_emails=max_emails)
        records: List[EmailRecord] = []
        progress = st.progress(0, text="Lettura email in corso...")
        total = len(ids) or 1
        for i, msg_id in enumerate(ids, start=1):
            try:
                status, data = mail.fetch(msg_id, "(RFC822)")
                if status != "OK" or not data:
                    continue
                raw = b""
                for item in data:
                    if isinstance(item, tuple):
                        raw += item[1]
                if not raw:
                    continue
                record = parse_email_message(index=len(records), uid=msg_id.decode(errors="ignore"), raw=raw, query_company=query_company)
                if record.match_score >= min_score or record.company_confidence >= 50:
                    record.index = len(records)
                    records.append(record)
            except Exception as exc:
                st.warning(f"Email saltata per errore lettura: {exc}")
            progress.progress(min(100, int(i * 100 / total)), text=f"Analizzate {i}/{total} email")
        progress.empty()
        return records
    finally:
        try:
            mail.close()
        except Exception:
            pass
        try:
            mail.logout()
        except Exception:
            pass


# -----------------------------
# Database e archiviazione
# -----------------------------

def db_path_for_archive(root: Path) -> Path:
    return root / "financeplus_email_archive.db"


def init_db(root: Path) -> sqlite3.Connection:
    root.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(db_path_for_archive(root))
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS archived_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT,
            item_type TEXT,
            sha256 TEXT,
            source_uid TEXT,
            message_id TEXT,
            mail_date TEXT,
            sender_name TEXT,
            sender_email TEXT,
            company TEXT,
            subject TEXT,
            original_filename TEXT,
            saved_filename TEXT,
            saved_path TEXT,
            document_type TEXT,
            status TEXT
        )
        """
    )
    conn.execute("CREATE INDEX IF NOT EXISTS idx_hash ON archived_items(sha256)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_company ON archived_items(company)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_subject ON archived_items(subject)")
    conn.commit()
    return conn


def existing_hash_path(conn: sqlite3.Connection, digest: str) -> Optional[str]:
    cur = conn.execute("SELECT saved_path FROM archived_items WHERE sha256 = ? AND status = 'SALVATO' LIMIT 1", (digest,))
    row = cur.fetchone()
    return row[0] if row else None


def insert_archived_item(
    conn: sqlite3.Connection,
    record: EmailRecord,
    item_type: str,
    digest: str,
    company: str,
    original_filename: str,
    saved_filename: str,
    saved_path: str,
    document_type: str,
    status: str,
) -> None:
    conn.execute(
        """
        INSERT INTO archived_items (
            created_at, item_type, sha256, source_uid, message_id, mail_date,
            sender_name, sender_email, company, subject, original_filename,
            saved_filename, saved_path, document_type, status
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            dt.datetime.now().isoformat(timespec="seconds"),
            item_type,
            digest,
            record.uid,
            record.message_id,
            record.date.isoformat(timespec="seconds"),
            record.sender_name,
            record.sender_email,
            company,
            record.subject,
            original_filename,
            saved_filename,
            saved_path,
            document_type,
            status,
        ),
    )
    conn.commit()


def build_mail_pdf_bytes(record: EmailRecord, company: str) -> bytes:
    if SimpleDocTemplate is None:
        # Fallback: testo in bytes, il chiamante salvera' .txt se ReportLab non c'e'.
        txt = (
            f"MAIL - {date_label_italian(record.date)}\n"
            f"Mittente: {record.from_header}\n"
            f"Data: {record.date}\n"
            f"Oggetto: {record.subject}\n"
            f"Azienda: {company}\n\n"
            f"Contenuto:\n{record.body_text}\n\n"
            f"Allegati:\n" + "\n".join(a.filename for a in record.attachments)
        )
        return txt.encode("utf-8", errors="replace")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=1.5 * cm, rightMargin=1.5 * cm, topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="FPTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, textColor=colors.HexColor("#0B3558"), spaceAfter=10))
    styles.add(ParagraphStyle(name="FPBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="FPLabel", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9.5, leading=12, textColor=colors.HexColor("#0B3558")))
    story = []
    story.append(Paragraph("FinancePlus - Stampa PDF Email", styles["FPTitle"]))
    meta = [
        ["Mittente", html.escape(record.from_header or "")],
        ["Data ricezione", record.date.strftime("%d/%m/%Y %H:%M")],
        ["Oggetto", html.escape(record.subject or "")],
        ["Azienda archivio", html.escape(company)],
        ["Confidenza", f"{record.company_confidence}%"],
    ]
    table_data = [[Paragraph(f"<b>{k}</b>", styles["FPBody"]), Paragraph(v, styles["FPBody"])] for k, v in meta]
    table = Table(table_data, colWidths=[4 * cm, 12 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF1F8")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C7D3")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.35 * cm))
    story.append(Paragraph("Contenuto email", styles["FPLabel"]))
    body = html.escape(record.body_text or "Corpo email non disponibile.").replace("\n", "<br/>")
    # Limite ragionevole per evitare PDF enormi; il corpo originale resta nella mail.
    if len(body) > 50000:
        body = body[:50000] + "<br/><br/>[Contenuto troncato nel PDF operativo]"
    story.append(Paragraph(body, styles["FPBody"]))
    story.append(Spacer(1, 0.35 * cm))
    story.append(Paragraph("Lista allegati", styles["FPLabel"]))
    if record.attachments:
        att_rows = [["Nome file", "Tipo", "Dimensione", "Classificazione"]]
        for a in record.attachments:
            att_rows.append([a.filename, a.content_type, f"{a.size:,}".replace(",", ".") + " byte", a.document_type])
        att_table = Table(att_rows, colWidths=[7 * cm, 3.2 * cm, 2.6 * cm, 3.2 * cm], repeatRows=1)
        att_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3558")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C7D3")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(att_table)
    else:
        story.append(Paragraph("Nessun allegato presente.", styles["FPBody"]))
    doc.build(story)
    return buffer.getvalue()


def archive_email_record(
    root: Path,
    record: EmailRecord,
    query_company: str,
    save_mail_pdf: bool = True,
    ignore_images: bool = True,
    force_query_company: bool = False,
    selected_attachment_indexes: Optional[Iterable[int]] = None,
    include_all_attachments: bool = True,
) -> Dict[str, object]:
    conn = init_db(root)
    selected_set = None if selected_attachment_indexes is None else set(selected_attachment_indexes)

    company = normalize_company_display(query_company) if force_query_company and query_company else record.detected_company
    if not company or company == "DA_VEDERE" or record.company_confidence < 35:
        company = normalize_company_display(query_company) if query_company else "DA_VEDERE"
    company_folder = sanitize_folder(company, default="DA_VEDERE")
    sender_folder = sanitize_folder(record.sender_name or record.sender_email or "MITTENTE_SCONOSCIUTO")
    dest = root / sender_folder / company_folder
    dest.mkdir(parents=True, exist_ok=True)

    saved = []
    duplicates = []
    skipped_images = []
    errors = []

    if save_mail_pdf:
        try:
            mail_bytes = build_mail_pdf_bytes(record, company)
            mail_ext = ".pdf" if SimpleDocTemplate is not None else ".txt"
            mail_name = f"MAIL_{date_label_italian(record.date)}{mail_ext}"
            digest = sha256_bytes(mail_bytes)
            existing = existing_hash_path(conn, digest)
            if existing:
                duplicates.append({"file": mail_name, "existing": existing})
                insert_archived_item(conn, record, "MAIL_PDF", digest, company, mail_name, mail_name, existing, "EMAIL", "DOPPIONE")
            else:
                path = unique_path(dest / mail_name)
                path.write_bytes(mail_bytes)
                insert_archived_item(conn, record, "MAIL_PDF", digest, company, mail_name, path.name, str(path), "EMAIL", "SALVATO")
                saved.append(str(path))
        except Exception as exc:
            errors.append(f"Errore salvataggio PDF mail {record.subject}: {exc}")

    for a in record.attachments:
        if selected_set is not None and a.index not in selected_set:
            continue
        if not include_all_attachments and selected_set is None:
            continue
        if ignore_images and is_image_file(a.filename):
            skipped_images.append(a.filename)
            insert_archived_item(conn, record, "ALLEGATO", a.sha256, company, a.filename, "", "", a.document_type, "IMMAGINE_IGNORATA")
            continue
        try:
            existing = existing_hash_path(conn, a.sha256)
            new_name = append_mail_date_to_filename(a.filename, record.date)
            if existing:
                duplicates.append({"file": a.filename, "existing": existing})
                insert_archived_item(conn, record, "ALLEGATO", a.sha256, company, a.filename, new_name, existing, a.document_type, "DOPPIONE")
                continue
            path = unique_path(dest / new_name)
            path.write_bytes(a.data)
            insert_archived_item(conn, record, "ALLEGATO", a.sha256, company, a.filename, path.name, str(path), a.document_type, "SALVATO")
            saved.append(str(path))
        except Exception as exc:
            errors.append(f"Errore salvataggio allegato {a.filename}: {exc}")
    return {"saved": saved, "duplicates": duplicates, "skipped_images": skipped_images, "errors": errors, "folder": str(dest), "company": company}


def get_archived_items(root: Path, limit: int = 500) -> List[Dict[str, object]]:
    db = db_path_for_archive(root)
    if not db.exists():
        return []
    conn = sqlite3.connect(db)
    conn.row_factory = sqlite3.Row
    cur = conn.execute(
        """
        SELECT created_at, item_type, mail_date, sender_name, sender_email, company, subject,
               original_filename, saved_filename, saved_path, document_type, status
        FROM archived_items
        ORDER BY id DESC
        LIMIT ?
        """,
        (limit,),
    )
    return [dict(r) for r in cur.fetchall()]


# -----------------------------
# Anteprime Streamlit
# -----------------------------

def preview_attachment(att: AttachmentData) -> None:
    filename = att.filename
    ext = Path(filename).suffix.lower()
    st.caption(f"{filename} - {att.content_type} - {att.size:,} byte".replace(",", "."))
    if ext == ".pdf":
        b64 = base64.b64encode(att.data).decode("ascii")
        components.html(
            f"""
            <div style="border:1px solid #ccd6df; height:720px; overflow:auto; background:#f8fafc;">
                <iframe src="data:application/pdf;base64,{b64}" width="100%" height="700px" style="border:0;"></iframe>
            </div>
            """,
            height=740,
        )
    elif is_image_file(filename):
        st.image(att.data, caption=filename, use_container_width=True)
    elif ext == ".docx" or ext == ".xlsx" or ext in PREVIEW_TEXT_EXTENSIONS:
        text = att.extracted_text or extract_text_from_attachment(filename, att.data)
        st.text_area("Anteprima testo", text or "Anteprima testuale non disponibile.", height=520, key=f"preview_{att.sha256}")
    else:
        st.info("Anteprima non disponibile per questo formato. Puoi comunque scaricarlo/archiviarlo.")
    st.download_button("Scarica allegato in locale", data=att.data, file_name=sanitize_filename(filename), mime=att.content_type or "application/octet-stream", key=f"download_{att.sha256}")


def record_label(record: EmailRecord) -> str:
    return f"{record.date.strftime('%d/%m/%Y')} | {record.sender_name} | {record.subject[:80]} | {len(record.attachments)} allegati"


# -----------------------------
# Interfaccia Streamlit
# -----------------------------

def setup_page() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="📧", layout="wide")
    st.markdown(
        """
        <style>
            .main-header {
                background: linear-gradient(90deg,#0B3558,#0E4A7B);
                color: white;
                padding: 18px 22px;
                border-radius: 14px;
                margin-bottom: 14px;
                box-shadow: 0 8px 20px rgba(11,53,88,.18);
            }
            .main-header h1 { margin: 0; font-size: 28px; }
            .main-header p { margin: 4px 0 0 0; color: #eaf1f8; }
            .fp-card {
                background: white;
                border: 1px solid #d9e3ec;
                border-radius: 12px;
                padding: 14px;
                box-shadow: 0 4px 12px rgba(11,53,88,.08);
            }
            .fp-warning { color:#8a4b00; font-weight:600; }
            .fp-ok { color:#0a6b3d; font-weight:600; }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(
        f"""
        <div class="main-header">
            <h1>{APP_TITLE}</h1>
            <p>Comando CERCA AZIENDA - Scarica tutto, vedi tutto, anteprima, sintesi e archivio per mittente/azienda.</p>
            <p>Versione {APP_VERSION}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def sidebar_config() -> Dict[str, object]:
    st.sidebar.header("Configurazione email")
    provider = st.sidebar.selectbox("Provider rapido", ["Gmail", "Outlook/Office365", "Aruba", "Libero", "Personalizzato"], index=0)
    defaults = {
        "Gmail": ("imap.gmail.com", 993),
        "Outlook/Office365": ("outlook.office365.com", 993),
        "Aruba": ("imaps.aruba.it", 993),
        "Libero": ("imapmail.libero.it", 993),
        "Personalizzato": ("imap.gmail.com", 993),
    }
    host_default, port_default = defaults[provider]
    host = st.sidebar.text_input("Server IMAP", host_default)
    port = st.sidebar.number_input("Porta", min_value=1, max_value=65535, value=port_default)
    username = st.sidebar.text_input("Email / username")
    password = st.sidebar.text_input("Password o App Password", type="password")
    mailbox = st.sidebar.text_input("Cartella email", "INBOX")
    use_ssl = st.sidebar.checkbox("Usa SSL", value=True)

    st.sidebar.header("Archivio")
    archive_root = st.sidebar.text_input("Cartella archivio locale", DEFAULT_ARCHIVE)
    ignore_images = st.sidebar.checkbox("Ignora immagini/loghi automatici", value=True, help="Consigliato: evita di salvare loghi e immagini decorative presenti nelle firme email.")
    force_query_company = st.sidebar.checkbox("Forza azienda cercata come cartella cliente", value=False, help="Se attivo, tutto cio' che scarichi viene salvato nella cartella dell'azienda scritta nella ricerca.")
    save_mail_pdf_selected = st.sidebar.checkbox("Salva sempre PDF della mail", value=True)

    st.sidebar.header("Ricerca")
    days_back = st.sidebar.number_input("Giorni indietro da analizzare", min_value=1, max_value=3650, value=365)
    max_emails = st.sidebar.number_input("Massimo email da leggere", min_value=10, max_value=5000, value=400, step=50)
    min_score = st.sidebar.slider("Soglia minima pertinenza", 0, 100, 25)

    st.sidebar.header("Sintesi AI opzionale")
    api_key = st.sidebar.text_input("OpenAI API key opzionale", type="password", help="Lascia vuoto per sintesi locale gratuita.")
    ai_model = st.sidebar.text_input("Modello AI", "gpt-4o-mini")

    return {
        "provider": provider,
        "host": host,
        "port": int(port),
        "username": username,
        "password": password,
        "mailbox": mailbox,
        "use_ssl": use_ssl,
        "archive_root": Path(archive_root).expanduser(),
        "ignore_images": ignore_images,
        "force_query_company": force_query_company,
        "save_mail_pdf_selected": save_mail_pdf_selected,
        "days_back": int(days_back),
        "max_emails": int(max_emails),
        "min_score": int(min_score),
        "api_key": api_key,
        "ai_model": ai_model,
    }


def page_cerca_azienda(config: Dict[str, object]) -> None:
    st.subheader("🔎 CERCA AZIENDA")
    st.write("Scrivi il nome dell'azienda. Il programma analizza oggetto, contenuto della mail, nomi allegati e testo estratto dagli allegati.")
    col_a, col_b, col_c = st.columns([2, 1, 1])
    with col_a:
        azienda = st.text_input("Nome azienda da cercare", placeholder="Esempio: BelGarden, SCHIANO S.R.L., PELCOM...")
    with col_b:
        search_button = st.button("🔎 Cerca", use_container_width=True, type="primary")
    with col_c:
        clear_button = st.button("Pulisci risultati", use_container_width=True)

    if clear_button:
        st.session_state.pop("email_results", None)
        st.session_state.pop("last_query_company", None)
        st.rerun()

    if search_button:
        if not azienda.strip():
            st.error("Inserisci il nome azienda da cercare.")
            return
        if not config["username"] or not config["password"]:
            st.error("Inserisci credenziali IMAP nella barra laterale.")
            return
        with st.spinner("Connessione alla casella email e analisi intelligente in corso..."):
            try:
                results = search_company_emails(
                    host=str(config["host"]),
                    port=int(config["port"]),
                    username=str(config["username"]),
                    password=str(config["password"]),
                    mailbox=str(config["mailbox"]),
                    query_company=azienda,
                    days_back=int(config["days_back"]),
                    max_emails=int(config["max_emails"]),
                    min_score=int(config["min_score"]),
                    use_ssl=bool(config["use_ssl"]),
                )
                st.session_state.email_results = results
                st.session_state.last_query_company = azienda
                st.success(f"Analisi completata: trovate {len(results)} email pertinenti.")
            except Exception as exc:
                st.error(f"Errore ricerca email: {exc}")

    results: List[EmailRecord] = st.session_state.get("email_results", [])
    query_company = st.session_state.get("last_query_company", azienda)
    if not results:
        st.info("Nessun risultato ancora caricato. Esegui una ricerca azienda.")
        return

    total_attachments = sum(len(r.attachments) for r in results)
    high_conf = sum(1 for r in results if r.company_confidence >= 75)
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Email trovate", len(results))
    col2.metric("Allegati", total_attachments)
    col3.metric("Alta confidenza", high_conf)
    col4.metric("Aziende rilevate", len(set(r.detected_company for r in results)))

    rows = []
    for r in results:
        rows.append({
            "#": r.index,
            "Data": r.date.strftime("%d/%m/%Y %H:%M"),
            "Mittente": r.sender_name,
            "Oggetto": r.subject,
            "Azienda rilevata": r.detected_company,
            "Confidenza": r.company_confidence,
            "Pertinenza": r.match_score,
            "Allegati": len(r.attachments),
        })
    st.dataframe(rows, use_container_width=True, hide_index=True)

    col_all, col_view = st.columns([1, 1])
    with col_all:
        if st.button("⬇️ SCARICA TUTTO", use_container_width=True, type="primary"):
            archive_root = Path(config["archive_root"])
            all_reports = []
            with st.spinner("Scarico e archiviazione di tutte le email/allegati..."):
                for r in results:
                    report = archive_email_record(
                        root=archive_root,
                        record=r,
                        query_company=query_company,
                        save_mail_pdf=True,
                        ignore_images=bool(config["ignore_images"]),
                        force_query_company=bool(config["force_query_company"]),
                        selected_attachment_indexes=None,
                        include_all_attachments=True,
                    )
                    all_reports.append(report)
            saved_count = sum(len(x["saved"]) for x in all_reports)
            duplicate_count = sum(len(x["duplicates"]) for x in all_reports)
            skipped_img_count = sum(len(x["skipped_images"]) for x in all_reports)
            st.success(f"Completato. Salvati {saved_count} file. Doppioni {duplicate_count}. Immagini ignorate {skipped_img_count}.")
            st.code(str(archive_root))
    with col_view:
        st.toggle("👁️ VEDI TUTTO - mostra anteprime, sintesi e selezione", key="show_all_results", value=True)

    if st.session_state.get("show_all_results", True):
        st.markdown("---")
        st.subheader("👁️ VEDI TUTTO")
        options: Dict[str, str] = {}
        for r in results:
            options[f"EMAIL::{r.index}"] = "EMAIL COMPLETA - " + record_label(r)
            for a in r.attachments:
                options[f"ATT::{r.index}::{a.index}"] = f"ALLEGATO - {r.date.strftime('%d/%m/%Y')} - {r.sender_name} - {a.filename}"
        selected_keys = st.multiselect("Seleziona email e/o allegati da scaricare", list(options.keys()), format_func=lambda k: options[k])
        if st.button("⬇️ SCARICA SELEZIONATI", use_container_width=True):
            if not selected_keys:
                st.warning("Seleziona almeno una email o un allegato.")
            else:
                keys_by_email: Dict[int, Dict[str, object]] = defaultdict(lambda: {"whole": False, "attachments": set()})
                for key in selected_keys:
                    parts = key.split("::")
                    if parts[0] == "EMAIL":
                        keys_by_email[int(parts[1])]["whole"] = True
                    elif parts[0] == "ATT":
                        keys_by_email[int(parts[1])]["attachments"].add(int(parts[2]))
                all_reports = []
                archive_root = Path(config["archive_root"])
                for idx, selection in keys_by_email.items():
                    record = next((r for r in results if r.index == idx), None)
                    if not record:
                        continue
                    if selection["whole"]:
                        report = archive_email_record(
                            archive_root, record, query_company=query_company,
                            save_mail_pdf=bool(config["save_mail_pdf_selected"]),
                            ignore_images=bool(config["ignore_images"]),
                            force_query_company=bool(config["force_query_company"]),
                            selected_attachment_indexes=None,
                            include_all_attachments=True,
                        )
                    else:
                        report = archive_email_record(
                            archive_root, record, query_company=query_company,
                            save_mail_pdf=bool(config["save_mail_pdf_selected"]),
                            ignore_images=bool(config["ignore_images"]),
                            force_query_company=bool(config["force_query_company"]),
                            selected_attachment_indexes=selection["attachments"],
                            include_all_attachments=False,
                        )
                    all_reports.append(report)
                saved_count = sum(len(x["saved"]) for x in all_reports)
                duplicate_count = sum(len(x["duplicates"]) for x in all_reports)
                skipped_img_count = sum(len(x["skipped_images"]) for x in all_reports)
                st.success(f"Selezione archiviata. Salvati {saved_count} file. Doppioni {duplicate_count}. Immagini ignorate {skipped_img_count}.")

        for r in results:
            with st.expander(record_label(r), expanded=False):
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Azienda", r.detected_company)
                c2.metric("Confidenza", f"{r.company_confidence}%")
                c3.metric("Pertinenza", f"{r.match_score}%")
                c4.metric("Allegati", len(r.attachments))
                st.write(f"**Mittente:** {r.from_header}")
                st.write(f"**Data:** {r.date.strftime('%d/%m/%Y %H:%M')}")
                st.write(f"**Oggetto:** {r.subject}")

                summary = ai_summary_optional(r, api_key=str(config.get("api_key") or ""), model=str(config.get("ai_model") or "gpt-4o-mini"))
                st.text_area("Sintesi intelligente", summary, height=170, key=f"summary_{r.uid}")
                st.text_area("Contenuto mail", r.body_text or "Corpo email non disponibile.", height=260, key=f"body_{r.uid}")

                if r.attachments:
                    st.markdown("**Allegati**")
                    att_names = [f"{a.index} - {a.filename} [{a.document_type}]" for a in r.attachments]
                    chosen = st.selectbox("Scegli allegato per anteprima", att_names, key=f"att_choice_{r.uid}")
                    chosen_index = int(chosen.split(" - ")[0])
                    att = next(a for a in r.attachments if a.index == chosen_index)
                    preview_attachment(att)
                else:
                    st.info("Questa email non contiene allegati.")


def page_archivio(config: Dict[str, object]) -> None:
    st.subheader("📁 Archivio e storico scarichi")
    root = Path(config["archive_root"])
    st.write("Cartella archivio:")
    st.code(str(root))
    items = get_archived_items(root)
    if not items:
        st.info("Nessuno storico disponibile. Dopo uno scarico verranno mostrati qui i file salvati, i doppioni e le immagini ignorate.")
        return
    st.dataframe(items, use_container_width=True, hide_index=True)
    csv = "\n".join([";".join(map(str, row.values())) for row in items])
    st.download_button("Scarica storico CSV", csv.encode("utf-8"), file_name="storico_archivio_email.csv", mime="text/csv")


def page_guida() -> None:
    st.subheader("ℹ️ Guida rapida")
    st.markdown(
        """
        ### Flusso operativo
        1. Compila le credenziali IMAP nella barra laterale.
        2. Scrivi il nome azienda in **CERCA AZIENDA**.
        3. Premi **Cerca**.
        4. Premi **VEDI TUTTO** per verificare email, allegati, anteprima e sintesi.
        5. Premi **SCARICA TUTTO** oppure seleziona solo cio' che serve e premi **SCARICA SELEZIONATI**.

        ### Archivio creato
        ```text
        Archivio_Email_Aziende/
          MITTENTE/
            AZIENDA/
              MAIL_06_MAGGIO_2026.pdf
              bilancio2025_06_MAGGIO_2026.pdf
        ```

        ### Regola importante
        Il programma usa prima **oggetto** e **contenuto della mail**. Se l'oggetto contiene
        "documentazione BelGarden", la mail e gli allegati vengono associati a BelGarden, salvo tua forzatura manuale.

        ### Gmail
        Per Gmail serve una **App Password** se hai la verifica in due passaggi. Non usare la password principale dell'account.
        """
    )


def main() -> None:
    setup_page()
    config = sidebar_config()
    tab1, tab2, tab3 = st.tabs(["🔎 CERCA AZIENDA", "📁 Archivio", "ℹ️ Guida rapida"])
    with tab1:
        page_cerca_azienda(config)
    with tab2:
        page_archivio(config)
    with tab3:
        page_guida()


if __name__ == "__main__":
    main()
